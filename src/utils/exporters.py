"""
ATS Resume Studio - Export Utilities
Generate downloadable files from text content.
"""

from __future__ import annotations
import io
from datetime import datetime


def text_to_docx_bytes(content: str, title: str = "Resume") -> bytes:
    """Convert text content to a formatted DOCX file in memory."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        lines = content.split("\n")
        first_line = True

        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue

            # Heuristic: all-caps short lines are headings
            is_heading = (
                line.isupper() and len(line) < 50 and not line.startswith("-")
            ) or (
                line.endswith(":") and len(line.split()) <= 4
            )

            if first_line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(20)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_line = False
            elif is_heading:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x6B, 0x46, 0xC1)
                # Add bottom border via paragraph format
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
            elif line.startswith("•") or line.startswith("-"):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(line.lstrip("•- "))
                p.paragraph_format.space_after = Pt(2)
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(2)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        # Fallback: return plain bytes if python-docx not installed
        return content.encode("utf-8")


def get_timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def create_download_filename(base: str, ext: str) -> str:
    return f"{base}_{get_timestamp()}.{ext}"
