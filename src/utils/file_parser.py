"""
ATS Resume Studio - File Parser Utilities
Handles PDF, DOCX, and TXT file extraction.
"""

from __future__ import annotations
import io
from typing import Optional


def extract_text_from_file(uploaded_file) -> tuple[str, str]:
    """
    Extract text from an uploaded Streamlit file object.
    Returns (extracted_text, error_message). error_message is empty on success.
    """
    if uploaded_file is None:
        return "", "No file provided."

    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()

    try:
        if filename.endswith(".txt"):
            return _extract_txt(file_bytes), ""
        elif filename.endswith(".pdf"):
            return _extract_pdf(file_bytes)
        elif filename.endswith(".docx"):
            return _extract_docx(file_bytes)
        else:
            return "", f"Unsupported file type: {filename.split('.')[-1].upper()}. Please upload PDF, DOCX, or TXT."
    except Exception as e:
        return "", f"Failed to parse file: {str(e)}"


def _extract_txt(file_bytes: bytes) -> str:
    """Decode plain text file."""
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _extract_pdf(file_bytes: bytes) -> tuple[str, str]:
    """Extract text from PDF using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        extracted = "\n".join(text_parts).strip()
        if not extracted:
            return "", "Could not extract text from PDF. Try copying the text manually."
        return extracted, ""
    except ImportError:
        return "", "pypdf not installed. Run: pip install pypdf"
    except Exception as e:
        return "", f"PDF parse error: {str(e)}"


def _extract_docx(file_bytes: bytes) -> tuple[str, str]:
    """Extract text from DOCX using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        extracted = "\n".join(paragraphs).strip()
        if not extracted:
            return "", "Could not extract text from DOCX file."
        return extracted, ""
    except ImportError:
        return "", "python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return "", f"DOCX parse error: {str(e)}"


def clean_text(text: str) -> str:
    """Basic text cleanup."""
    import re
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()
